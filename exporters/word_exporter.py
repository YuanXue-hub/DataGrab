"""Word 文档导出器

将爬取的数据导出为格式化的 .docx 文件。
支持新闻、军事、经济、社交媒体等多种数据类型的排版。
"""

from typing import List, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from loguru import logger

from storage.models import (
    DataItem, NewsArticle, MilitaryData, EconomicData, SocialPost,
)
from storage.repository import Repository


class WordExporter:
    """Word 文档导出器

    将 Repository 中的数据生成格式化的 Word 报告。
    """

    def __init__(
        self,
        report_title: str = "俄乌冲突数据报告",
        report_subtitle: str = "Russia-Ukraine Conflict Data Report",
        group_by: str = "source",  # source | date | type
    ):
        """
        Args:
            report_title: 报告中文标题
            report_subtitle: 报告英文副标题
            group_by: 分组方式
        """
        self.report_title = report_title
        self.report_subtitle = report_subtitle
        self.group_by = group_by
        self.doc: Optional[Document] = None

    def export(
        self,
        repository: Repository,
        output_path: str,
    ) -> str:
        """导出数据到 Word 文档

        Args:
            repository: 数据仓库
            output_path: 输出文件路径

        Returns:
            输出文件的绝对路径
        """
        self.doc = Document()
        self._setup_styles()

        # 封面标题
        self._write_title_page()

        # 摘要统计
        self._write_summary(repository)

        # 按来源分组输出内容
        self._write_content_by_source(repository)

        # 保存文档
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        # 添加时间戳到文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{output.stem}_{timestamp}{output.suffix}"
        final_path = output.parent / filename

        self.doc.save(str(final_path))
        logger.info(f"Word document saved to: {final_path}")
        return str(final_path)

    def _setup_styles(self):
        """配置文档样式"""
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        # 设置中文字体回退
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _write_title_page(self):
        """写入封面标题页"""
        # 空行间距
        for _ in range(4):
            self.doc.add_paragraph()

        # 中文标题
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.report_title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # 英文副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(self.report_subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 生成时间
        self.doc.add_paragraph()
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 数据来源说明
        self.doc.add_paragraph()
        note = self.doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run("本报告数据来自公开网络来源，仅供研究参考")
        run.font.size = Pt(9)
        run.font.italic = True

        self.doc.add_page_break()

    def _write_summary(self, repository: Repository):
        """写入摘要统计"""
        summary = repository.summary()

        # 摘要标题
        heading = self.doc.add_heading("数据摘要 / Data Summary", level=1)

        # 统计表格
        table = self.doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "指标 / Metric"
        hdr_cells[1].text = "数量 / Count"

        # 数据行
        rows_data = [
            ("总数据条数 / Total Items", str(summary["total"])),
            ("新闻文章 / News Articles", str(summary["news"])),
            ("军事数据 / Military Data", str(summary["military"])),
            ("经济数据 / Economic Data", str(summary["economic"])),
            ("社交媒体 / Social Media Posts", str(summary["social"])),
        ]

        for label, value in rows_data:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        # 来源明细
        if summary.get("by_source"):
            self.doc.add_paragraph()
            source_heading = self.doc.add_heading("数据来源明细 / Source Breakdown", level=2)
            source_table = self.doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
            hdr = source_table.rows[0].cells
            hdr[0].text = "数据源 / Source"
            hdr[1].text = "数量 / Count"

            for source, count in summary["by_source"].items():
                row = source_table.add_row()
                row.cells[0].text = source
                row.cells[1].text = str(count)

        self.doc.add_page_break()

    def _write_content_by_source(self, repository: Repository):
        """按数据源分组输出内容"""
        # 获取所有数据源
        items = repository.get_all()
        by_source = {}
        for item in items:
            source = getattr(item, "source_name", "unknown")
            by_source.setdefault(source, []).append(item)

        for source, source_items in by_source.items():
            self._write_source_section(source, source_items)

    def _write_source_section(self, source_name: str, items: List[DataItem]):
        """写入单个数据源的内容区域"""
        # 数据源标题
        heading = self.doc.add_heading(
            f"数据源: {source_name.upper()}",
            level=1,
        )

        # 分类处理
        news_items = [i for i in items if isinstance(i, NewsArticle)]
        military_items = [i for i in items if isinstance(i, MilitaryData)]
        social_items = [i for i in items if isinstance(i, SocialPost)]
        economic_items = [i for i in items if isinstance(i, EconomicData)]

        # 新闻文章
        if news_items:
            self.doc.add_heading(f"新闻文章 ({len(news_items)}篇)", level=2)
            for article in news_items[:30]:  # 每个源最多30篇
                self._write_article(article)

        # 军事数据
        if military_items:
            self.doc.add_heading(f"军事数据 ({len(military_items)}条)", level=2)
            self._write_military_table(military_items)

        # 社交媒体
        if social_items:
            self.doc.add_heading(f"社交媒体 ({len(social_items)}条)", level=2)
            for post in social_items[:20]:
                self._write_social_post(post)

        # 经济数据
        if economic_items:
            self.doc.add_heading(f"经济数据 ({len(economic_items)}条)", level=2)
            for econ in economic_items[:20]:
                self._write_economic_data(econ)

        self.doc.add_page_break()

    def _write_article(self, article: NewsArticle):
        """写入单篇新闻文章"""
        # 标题
        title_para = self.doc.add_paragraph()
        run = title_para.add_run(article.title)
        run.font.bold = True
        run.font.size = Pt(13)

        # 元信息
        meta_parts = []
        if article.published_at:
            meta_parts.append(f"📅 {article.published_at.strftime('%Y-%m-%d %H:%M')}")
        meta_parts.append(f"🌐 {article.language.upper()}")
        if article.tags:
            meta_parts.append(f"🏷 {', '.join(article.tags[:5])}")

        meta_para = self.doc.add_paragraph()
        run = meta_para.add_run(" | ".join(meta_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # 摘要
        if article.summary:
            summary_text = article.summary[:300]
            summary_para = self.doc.add_paragraph()
            run = summary_para.add_run(f"摘要: {summary_text}")
            run.font.size = Pt(10)
            run.font.italic = True

        # 正文（截取前500字）
        if article.content:
            content_text = article.content[:500]
            if len(article.content) > 500:
                content_text += "..."
            content_para = self.doc.add_paragraph()
            run = content_para.add_run(content_text)
            run.font.size = Pt(10)

        # 原文链接
        if article.source_url:
            link_para = self.doc.add_paragraph()
            run = link_para.add_run(f"原文链接: {article.source_url}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x33, 0x66, 0xcc)

        # 分隔线
        self.doc.add_paragraph("─" * 50)

    def _write_military_table(self, items: List[MilitaryData]):
        """将军事数据写入表格"""
        if not items:
            return

        table = self.doc.add_table(rows=1, cols=8, style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        headers = ["数据类型", "指标", "数值", "单位", "阵营", "语言", "置信度", "时间"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # 表头加粗
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        for item in items[:50]:
            row = table.add_row()
            cells = row.cells
            cells[0].text = item.data_type or "-"
            cells[1].text = item.metric_name or "-"
            cells[2].text = str(item.value) if item.value is not None else "-"
            cells[3].text = item.unit or "-"
            cells[4].text = item.side or "-"
            cells[5].text = getattr(item, "language", "en").upper()
            cells[6].text = item.confidence or "-"
            reported = item.reported_at.strftime("%Y-%m-%d") if item.reported_at else "-"
            cells[7].text = reported

    def _write_social_post(self, post: SocialPost):
        """写入单条社交媒体帖子"""
        # 作者和平台
        header_para = self.doc.add_paragraph()
        run = header_para.add_run(f"@{post.author}")
        run.font.bold = True
        header_para.add_run(f"  [{post.platform.upper()}]")
        lang_code = getattr(post, "language", "en").upper()
        header_para.add_run(f"  🌐 {lang_code}")

        # 内容
        if post.content:
            content_para = self.doc.add_paragraph()
            run = content_para.add_run(post.content[:300])
            run.font.size = Pt(10)

        # 互动数据
        if post.engagement:
            eng_parts = []
            for k, v in post.engagement.items():
                eng_parts.append(f"{k}: {v}")
            eng_para = self.doc.add_paragraph()
            run = eng_para.add_run(" | ".join(eng_parts))
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # 发布时间
        if post.posted_at:
            time_para = self.doc.add_paragraph()
            run = time_para.add_run(post.posted_at.strftime("%Y-%m-%d %H:%M"))
            run.font.size = Pt(8)

        self.doc.add_paragraph("· · ·")

    def _write_economic_data(self, item: EconomicData):
        """写入经济数据"""
        para = self.doc.add_paragraph()
        lang = getattr(item, "language", "en").upper()
        run = para.add_run(f"📊 {item.indicator}: {item.description}  🌐 {lang}")
        run.font.bold = True

        if item.value is not None:
            value_para = self.doc.add_paragraph()
            value_para.add_run(f"  值: {item.value} {item.unit}")

        if item.country:
            value_para = self.doc.add_paragraph()
            value_para.add_run(f"  国家/地区: {item.country}")

        if item.source_url:
            link_para = self.doc.add_paragraph()
            run = link_para.add_run(f"  来源: {item.source_url}")
            run.font.size = Pt(8)


def export_to_word(
    repository: Repository,
    output_path: str,
    report_title: str = "俄乌冲突数据报告",
) -> str:
    """便捷函数：将仓库数据导出为 Word 文档

    Args:
        repository: 数据仓库
        output_path: 输出文件路径
        report_title: 报告标题

    Returns:
        输出文件路径
    """
    exporter = WordExporter(report_title=report_title)
    return exporter.export(repository, output_path)
